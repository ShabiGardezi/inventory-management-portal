from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = "/mnt/data/Inventory_Management_System_Master_Documentation.docx"
FALLBACK_OUTPUT_PATH = os.path.join(
    os.path.dirname(__file__),
    "..",
    "mnt",
    "data",
    "Inventory_Management_System_Master_Documentation.docx",
)
SYSTEM_NAME = "Inventory Management System"
VERSION = "1.0"
PREPARED_BY = "Engineering / Inventory Management Portal Team"
TODAY = date.today().strftime("%Y-%m-%d")
CONFIDENTIAL_NOTE = "CONFIDENTIAL — Internal Use Only"


def _add_field_simple(paragraph, instr: str) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = " "
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Margins for printing
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # Header (default, not first page)
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = f"{SYSTEM_NAME} — v{VERSION}"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Footer (default, not first page): table for left/right layout
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""

    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]

    lp = left.paragraphs[0]
    lp.text = CONFIDENTIAL_NOTE
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.add_run("Page ")
    _add_field_simple(rp, "PAGE")
    rp.add_run(" of ")
    _add_field_simple(rp, "NUMPAGES")


def _title_page(doc: Document) -> None:
    # First-page header/footer intentionally left blank by different_first_page_header_footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SYSTEM_NAME)
    r.bold = True
    r.font.size = Pt(28)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Master Documentation (Phase 0 → Phase 5)")
    r2.font.size = Pt(16)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version: {VERSION}\n").bold = True
    meta.add_run(f"Prepared by: {PREPARED_BY}\n")
    meta.add_run(f"Date: {TODAY}\n")

    doc.add_page_break()


def _toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    _add_field_simple(p, 'TOC \\o \"1-3\" \\h \\z \\u')
    doc.add_page_break()


def _version_history(doc: Document) -> None:
    doc.add_heading("Version History", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Version"
    hdr[1].text = "Date"
    hdr[2].text = "Description"
    hdr[3].text = "Author"

    row = table.add_row().cells
    row[0].text = VERSION
    row[1].text = TODAY
    row[2].text = "Initial master documentation for Phase 0–5 implementation."
    row[3].text = PREPARED_BY

    doc.add_page_break()


def _h1(doc: Document, title: str) -> None:
    doc.add_heading(title, level=1)


def _h2(doc: Document, title: str) -> None:
    doc.add_heading(title, level=2)


def _h3(doc: Document, title: str) -> None:
    doc.add_heading(title, level=3)


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _num(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def build_doc() -> Document:
    doc = Document()
    _set_normal_style(doc)
    _add_header_footer(doc)

    _title_page(doc)
    _toc(doc)
    _version_history(doc)

    # ---------------------------------------------------------
    # SECTION 1 — EXECUTIVE SUMMARY
    # ---------------------------------------------------------
    _h1(doc, "SECTION 1 — Executive Summary")
    _h2(doc, "System Overview")
    _bullet(doc, "A modular inventory management platform designed for controlled stock execution, auditability, and operational scale.")
    _bullet(doc, "Supports warehouses, products, stock ledger, approvals, forecasting, and scanning workflows up to Phase 5.")
    _h2(doc, "Target Audience")
    _bullet(doc, "Operations (warehouse staff, inventory clerks, managers).")
    _bullet(doc, "Finance/controls (read-only financial visibility where enabled).")
    _bullet(doc, "Administrators and auditors (RBAC, audit logs, approvals).")
    _h2(doc, "Competitive Advantages")
    _bullet(doc, "Ledger-first stock engine: immutable movements with balance snapshots.")
    _bullet(doc, "Approval workflow to enforce governance before execution.")
    _bullet(doc, "Batch and serial tracking support for regulated/traceable inventory.")
    _bullet(doc, "Barcode/scan lookup for high-throughput operations.")
    _h2(doc, "Enterprise Readiness")
    _bullet(doc, "Strict RBAC permissions, system roles plus customizable roles.")
    _bullet(doc, "Audit logging for critical actions.")
    _bullet(doc, "Integrity verification scripts for ledger vs balances.")
    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 2 — SYSTEM ARCHITECTURE
    # ---------------------------------------------------------
    _h1(doc, "SECTION 2 — System Architecture")
    _h2(doc, "Technical Stack")
    _bullet(doc, "Next.js 15 (App Router), TypeScript (strict).")
    _bullet(doc, "Prisma ORM + PostgreSQL.")
    _bullet(doc, "NextAuth (Credentials) for authentication; JWT session strategy.")
    _h2(doc, "Service Structure (Architecture Rules)")
    _bullet(doc, "Route handlers are thin; business logic lives in server/services; data access in server/repositories.")
    _bullet(doc, "StockService is the only stock mutator; do not write stock ledger/balances directly.")
    _h2(doc, "Stock Engine: Ledger + Balances")
    _bullet(doc, "Ledger: stock_movements (immutable; IN/OUT) with references (PURCHASE/SALE/TRANSFER/ADJUSTMENT).")
    _bullet(doc, "Balances: stock_balances as snapshot per product/warehouse (+ optional batchId).")
    _bullet(doc, "Transfers execute as two movements with a shared referenceId in one transaction.")
    _h2(doc, "Approval Engine (Phase 4)")
    _bullet(doc, "Policies define which entity actions require approval.")
    _bullet(doc, "ApprovalRequest lifecycle: PENDING → APPROVED/REJECTED/CANCELLED.")
    _bullet(doc, "Execution is idempotent: approving an already-approved request does not execute twice.")
    _h2(doc, "Forecasting Engine (Phase 3)")
    _bullet(doc, "InventoryMetricsService computes avgDailySales, daysOfCover, predictedStockoutDate, suggestedReorderQty.")
    _bullet(doc, "ReorderPolicy per product+warehouse defines leadTimeDays, min/max, safety stock.")
    _h2(doc, "Scanning Architecture (Phase 5)")
    _bullet(doc, "Scan lookup endpoint resolves codes (e.g., product.barcode / sku / serial numbers) for quick retrieval.")
    _bullet(doc, "Database indexes support lookup performance (e.g., Product.barcode, ProductSerial.serialNumber).")
    _h2(doc, "High-Level Data Flow")
    _bullet(doc, "UI → API route handler → service layer (validation + rules) → repositories/Prisma → audit log.")
    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 3 — DATABASE & BUSINESS LOGIC
    # ---------------------------------------------------------
    _h1(doc, "SECTION 3 — Database & Business Logic")
    _h2(doc, "Ledger Principle")
    _bullet(doc, "All inventory changes are recorded as immutable stock_movements.")
    _bullet(doc, "Movements are never edited/deleted; corrections happen via new movements (e.g., adjustment).")
    _h2(doc, "Balance Snapshot Principle")
    _bullet(doc, "stock_balances stores the current snapshot for fast reads; it must match SUM(ledger movements).")
    _bullet(doc, "Integrity verification checks balances vs ledger and transfer consistency.")
    _h2(doc, "Transfers Create Two Movements")
    _bullet(doc, "A transfer creates one OUT movement from source and one IN movement to destination, with the same referenceId.")
    _h2(doc, "Approvals Block Execution")
    _bullet(doc, "When approval is required, the system creates a request and defers stock mutation until approved.")
    _h2(doc, "Batch & Serial Rules")
    _bullet(doc, "Batch-tracked products require batchId/batchNumber on IN and batchId on OUT.")
    _bullet(doc, "Serial-tracked products require serial numbers on OUT; serials transition status (e.g., IN_STOCK → SOLD).")
    _h2(doc, "FIFO vs Average Cost (Phase 2)")
    _bullet(doc, "Valuation method is stored in Settings. InventoryLayer/Consumption tables exist to support costing.")
    _bullet(doc, "COGS and margin fields exist on sales items for reporting when valuation/COGS logic is enabled.")
    _h2(doc, "Reorder Forecasting Formulas (Phase 3)")
    _bullet(doc, "Avg daily sales = total sold qty over lookback / lookbackDays.")
    _bullet(doc, "Days of cover = currentStock / avgDailySales (if avgDailySales > 0).")
    _bullet(doc, "Suggested reorder qty = max(0, leadDemand + safetyStock − currentStock).")
    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 4 — COMPLETE MODULE DOCUMENTATION
    # ---------------------------------------------------------
    _h1(doc, "SECTION 4 — Complete Module Documentation")

    modules = [
        "Dashboard",
        "Products",
        "Warehouses",
        "Stock Movements",
        "Purchases",
        "Sales",
        "Reports",
        "Users",
        "Roles",
        "Audit Logs",
        "Settings",
        "Approvals",
        "Scan",
    ]

    for m in modules:
        _h2(doc, m)
        _h3(doc, "Purpose")
        _bullet(doc, f"Provides the primary capabilities for {m.lower()} management and visibility.")
        _h3(doc, "What it manages")
        _bullet(doc, "Entities, validations, and operational workflows relevant to this module.")
        _h3(doc, "Core functionality")
        _bullet(doc, "Create, view, and manage records according to permissions.")
        _h3(doc, "Business rules")
        _bullet(doc, "Rules enforced by services (stock execution rules, approval gating, validation constraints).")
        _h3(doc, "Validation rules")
        _bullet(doc, "Schema validation (API payload validation) and service-level checks (e.g., stock availability).")
        _h3(doc, "Data flow")
        _bullet(doc, "UI → API → service → Prisma → audit/integrity metrics updates.")
        _h3(doc, "Role access overview")
        _bullet(doc, "Access is permission-based; Admin has full access; other roles are scoped.")
        _h3(doc, "Common use cases")
        _bullet(doc, "Daily operations, reporting, and exception handling.")
        _h3(doc, "Edge cases")
        _bullet(doc, "Approval-required actions, insufficient stock, batch/serial requirements, and invalid inputs.")

    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 5 — ROLE-BASED COMPLETE FUNCTIONAL GUIDE
    # ---------------------------------------------------------
    _h1(doc, "SECTION 5 — Role-Based Complete Functional Guide")
    _h2(doc, "Role Model")
    _bullet(doc, "System supports default roles (Admin/Manager/Staff/Viewer) and custom roles (permission-based).")
    _bullet(doc, "Custom roles can be created and assigned with specific permissions (e.g., warehouse_lead).")

    _h2(doc, "Admin")
    _bullet(doc, "Can see/create/edit/delete across modules; can configure approvals, roles, and settings.")
    _bullet(doc, "Approval authority depends on policy and permissions (approvals.review/manage).")
    _h3(doc, "Admin step-by-step guide")
    for step in [
        "Log in as admin and review dashboard health.",
        "Create warehouses and confirm codes/locations.",
        "Configure valuation method (FIFO or Average) and financial visibility rules.",
        "Configure approval policies for purchase receive, sale confirm, transfers, and adjustments.",
        "Create roles and users; assign permissions based on job function.",
        "Configure reorder policies per product/warehouse; recompute metrics.",
        "Review audit logs and approvals queue regularly.",
        "Enable/disable system lockdown according to governance policy.",
        "Run reports and export where permitted.",
    ]:
        _num(doc, step)

    _h2(doc, "Manager")
    _bullet(doc, "Can operate inventory workflows, review reports, and review/approve where permitted.")
    _h3(doc, "Manager step-by-step guide")
    for step in [
        "Monitor dashboard KPIs and low-stock indicators.",
        "Review the approvals queue and approve eligible requests.",
        "Oversee purchase receiving and investigate exceptions (batch/serial mismatches).",
        "Review stock movement history for anomalies.",
        "Use reports to support replenishment decisions.",
    ]:
        _num(doc, step)

    _h2(doc, "Staff")
    _bullet(doc, "Executes operational workflows (receive, confirm, transfer, adjust) subject to permissions and approvals.")
    _h3(doc, "Staff step-by-step guide")
    for step in [
        "Create or prepare purchase receive payload; include batch/serial inputs when required.",
        "Receive purchase (or submit for approval if enabled).",
        "Create sale and confirm sale; for serial items, select serial numbers; for batch items, select batch.",
        "Transfer stock between warehouses (may require approval).",
        "Perform stock adjustments with reason codes (may require approval).",
        "Use scanning via USB/camera for fast product lookup and operational flow entry.",
    ]:
        _num(doc, step)

    _h2(doc, "Viewer")
    _bullet(doc, "Read-only access to dashboards and reports; cannot execute stock-changing actions.")
    _h3(doc, "Viewer step-by-step guide")
    for step in [
        "View dashboard for high-level visibility.",
        "Use reports to review inventory and movement history.",
        "Export reports only if permission allows; otherwise request access from Admin.",
    ]:
        _num(doc, step)

    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 6 — COMPLETE FEATURE FLOWS
    # ---------------------------------------------------------
    _h1(doc, "SECTION 6 — Complete Feature Flows")

    flows = [
        "Purchase Receive Flow",
        "Sale Confirm Flow",
        "Transfer Flow",
        "Stock Adjustment Flow",
        "Approval Flow",
        "Batch Lifecycle",
        "Serial Lifecycle",
        "FIFO Layer Consumption Flow",
        "Average Cost Flow",
        "Reorder Forecast Flow",
        "Scan & Lookup Flow",
    ]

    for f in flows:
        _h2(doc, f)
        _h3(doc, "Trigger")
        _bullet(doc, "User action via UI/API route that initiates the workflow.")
        _h3(doc, "Validation")
        _bullet(doc, "Schema validation + business rule checks (permissions, quantities, batch/serial requirements).")
        _h3(doc, "Service execution")
        _bullet(doc, "Service layer executes stock mutation through StockService when allowed.")
        _h3(doc, "DB impact")
        _bullet(doc, "Writes ledger movements and updates balances; persists request entities where applicable.")
        _h3(doc, "Status transitions")
        _bullet(doc, "For approval workflows: PENDING_APPROVAL → APPROVED/REJECTED → APPLIED/CONFIRMED/RECEIVED.")
        _h3(doc, "Audit logs created")
        _bullet(doc, "Critical actions create audit entries; approvals add request/review/execution logs.")

    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 7 — SMART REORDER & FORECASTING
    # ---------------------------------------------------------
    _h1(doc, "SECTION 7 — Smart Reorder & Forecasting")
    _h2(doc, "Formulas")
    _bullet(doc, "Avg daily sales = SUM(sales qty over lookback) / lookbackDays.")
    _bullet(doc, "Days of cover = currentStock / avgDailySales (when avgDailySales > 0).")
    _bullet(doc, "Suggested reorder qty = max(0, leadTimeDays*avgDailySales + safetyStock − currentStock).")
    _h2(doc, "Stockout Prediction Logic")
    _bullet(doc, "predictedStockoutDate is computed as now + daysOfCover when avgDailySales > 0.")
    _h2(doc, "Dashboard Integration")
    _bullet(doc, "Metrics are stored in inventory_metrics for fast dashboard rendering.")
    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 8 — VALUATION & FINANCIAL LOGIC
    # ---------------------------------------------------------
    _h1(doc, "SECTION 8 — Valuation & Financial Logic")
    _h2(doc, "FIFO Example")
    _bullet(doc, "Example: buy 10 units @ $5, then 10 units @ $7. Sell 12 units → COGS = 10*$5 + 2*$7 = $64.")
    _h2(doc, "Average Cost Example")
    _bullet(doc, "Example: same purchases → average = (10*$5 + 10*$7)/20 = $6. Sell 12 → COGS = 12*$6 = $72.")
    _h2(doc, "COGS and Margin")
    _bullet(doc, "COGS is recorded per sale item when valuation logic is enabled; margin = revenue − COGS.")
    _h2(doc, "Permission Gating")
    _bullet(doc, "Financial visibility should be controlled by permissions (e.g., financials.read) and Settings.showFinancials.")
    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 9 — APPROVAL WORKFLOW ENGINE
    # ---------------------------------------------------------
    _h1(doc, "SECTION 9 — Approval Workflow Engine")
    _h2(doc, "Policy Configuration")
    _bullet(doc, "ApprovalPolicy per entity type toggles whether approval is required.")
    _bullet(doc, "Policies can include permission requirements and thresholds (e.g., minAmount).")
    _h2(doc, "Lifecycle: Pending → Approved → Executed")
    _bullet(doc, "Request is created PENDING; reviewers approve/reject; approval execution performs the stock mutation.")
    _h2(doc, "Idempotency")
    _bullet(doc, "If an approval is already approved/executed, re-approving will not execute again.")
    _h2(doc, "Rejection Behavior")
    _bullet(doc, "Rejected requests do not mutate stock; entity status transitions to REJECTED where applicable.")
    _h2(doc, "Audit Trail")
    _bullet(doc, "Approval requests, reviews, cancellations, and executions are audited.")
    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 10 — BARCODE & SCANNING GUIDE
    # ---------------------------------------------------------
    _h1(doc, "SECTION 10 — Barcode & Scanning Guide")
    _h2(doc, "Modes")
    _bullet(doc, "Manual entry: user types or pastes a code.")
    _bullet(doc, "USB scanner: behaves like keyboard input, submits the scanned code.")
    _bullet(doc, "Camera scanning: uses device camera; best on mobile over HTTPS.")
    _h2(doc, "Lookup Resolution Priority")
    _bullet(doc, "Typical resolution: product barcode → SKU → serial number → batch barcode/number (as supported).")
    _h2(doc, "Performance & Indexing")
    _bullet(doc, "Ensure indexed columns exist (Product.barcode, ProductSerial.serialNumber, Batch.barcode).")
    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 11 — ONBOARDING GUIDE FOR NEW BUSINESS
    # ---------------------------------------------------------
    _h1(doc, "SECTION 11 — Onboarding Guide for New Business")
    for step in [
        "Create warehouses and verify default warehouse settings.",
        "Import or create products (enable batch/serial flags where required).",
        "Perform initial stock load via purchase receive or adjustments (with correct batch/serial inputs).",
        "Configure approvals policies and reviewer roles (if governance required).",
        "Configure reorder policies and recompute metrics.",
        "Validate scanning setup (manual/USB/camera) and label process.",
        "Run integrity verification and confirm dashboards/reports are correct.",
        "Go-live checklist: restrict admin access, confirm backups, enable/disable system lockdown policy.",
    ]:
        _num(doc, step)
    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 12 — TROUBLESHOOTING & FAQ
    # ---------------------------------------------------------
    _h1(doc, "SECTION 12 — Troubleshooting & FAQ")
    _h2(doc, "Common Issues")
    _bullet(doc, "Stock not updating: ensure action was executed (not pending approval) and that StockService was used.")
    _bullet(doc, "Approval blocking execution: check ApprovalPolicy and approval request status.")
    _bullet(doc, "Serial mismatch: serialNumbers count must equal quantity; serial must be IN_STOCK in the selected warehouse.")
    _bullet(doc, "Batch required: batch-tracked products require batch input/selection for IN/OUT.")
    _bullet(doc, "Negative stock blocked: Settings.allowNegativeStock=false prevents stock from going below zero.")
    _bullet(doc, "Scan not found: verify code source (barcode vs SKU vs serial), indexing, and product activation.")
    _bullet(doc, "Dashboard delay: metrics recomputation may run after execution; recompute metrics if needed.")
    _bullet(doc, "Integrity mismatch: run integrity checks; investigate missing transfer pairs or manual data edits.")
    doc.add_page_break()

    # ---------------------------------------------------------
    # SECTION 13 — GLOSSARY
    # ---------------------------------------------------------
    _h1(doc, "SECTION 13 — Glossary")
    glossary = [
        ("Ledger", "Immutable record of stock movements (IN/OUT) used as the source of truth."),
        ("Balance", "Snapshot quantity stored for fast reads; must reconcile with the ledger."),
        ("Batch/Lot", "A grouped inventory identifier (e.g., manufacturing batch) used for traceability."),
        ("Serial", "A unique identifier per unit; tracked through lifecycle states."),
        ("FIFO", "First-In, First-Out valuation method consuming oldest layers first."),
        ("Average Cost", "Valuation method using weighted average unit cost."),
        ("COGS", "Cost of Goods Sold; the cost portion of a sale."),
        ("Layer", "Inventory valuation layer representing acquired quantity at a unit cost."),
        ("Days of Cover", "How long current stock lasts at the observed sales rate."),
        ("Approval Policy", "Rules defining when an action requires approval."),
        ("Approval Request", "A request instance awaiting review; approval triggers execution."),
    ]
    for term, desc in glossary:
        _h2(doc, term)
        doc.add_paragraph(desc)

    return doc


def main() -> None:
    doc = build_doc()
    try:
        os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
        doc.save(OUTPUT_PATH)
        print(f"Saved: {OUTPUT_PATH}")
        return
    except OSError:
        pass

    fallback_dir = os.path.dirname(os.path.abspath(FALLBACK_OUTPUT_PATH))
    os.makedirs(fallback_dir, exist_ok=True)
    doc.save(os.path.abspath(FALLBACK_OUTPUT_PATH))
    print(f"Saved (fallback): {os.path.abspath(FALLBACK_OUTPUT_PATH)}")


if __name__ == "__main__":
    main()

