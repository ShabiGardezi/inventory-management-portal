from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = "/mnt/data/Inventory_Management_System_Client_Proposal.docx"
FALLBACK_OUTPUT_PATH = os.path.join(
    os.path.dirname(__file__),
    "..",
    "mnt",
    "data",
    "Inventory_Management_System_Client_Proposal.docx",
)

SYSTEM_NAME = "Inventory Management System"
VERSION = "1.0"
TODAY = date.today().strftime("%Y-%m-%d")


def _add_field_simple(paragraph, instr: str) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = " "
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def _set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def _add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Print-ready margins
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)

    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = SYSTEM_NAME
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Page ")
    _add_field_simple(fp, "PAGE")
    fp.add_run(" of ")
    _add_field_simple(fp, "NUMPAGES")


def _cover_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SYSTEM_NAME)
    r.bold = True
    r.font.size = Pt(30)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Client Proposal")
    r2.font.size = Pt(16)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version: {VERSION}\n").bold = True
    meta.add_run(f"Date: {TODAY}\n")
    meta.add_run("Prepared for: ________________________________\n")
    meta.add_run("Prepared by: ________________________________\n")

    doc.add_page_break()


def _toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    _add_field_simple(p, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()


def _h1(doc: Document, title: str) -> None:
    doc.add_heading(title, level=1)


def _h2(doc: Document, title: str) -> None:
    doc.add_heading(title, level=2)


def _h3(doc: Document, title: str) -> None:
    doc.add_heading(title, level=3)


def _p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _num(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def _benefits_block(doc: Document, *, business: str, operational: str, risk: str) -> None:
    _h3(doc, "Benefits")
    _bullet(doc, f"Business benefit: {business}")
    _bullet(doc, f"Operational benefit: {operational}")
    _bullet(doc, f"Risk mitigation benefit: {risk}")


def _feature_section(
    doc: Document,
    *,
    title: str,
    overview: str,
    key_capabilities: list[str],
    business_benefit: str,
    operational_benefit: str,
    risk_benefit: str,
    operational_notes: list[str],
    kpis: list[str],
    typical_outcomes: list[str],
    example_steps: list[str],
) -> None:
    _h2(doc, title)
    _p(doc, overview)
    _h3(doc, "Key capabilities")
    for cap in key_capabilities:
        _bullet(doc, cap)
    _benefits_block(doc, business=business_benefit, operational=operational_benefit, risk=risk_benefit)
    _h3(doc, "Operational notes (enterprise-friendly)")
    for note in operational_notes:
        _bullet(doc, note)
    _h3(doc, "KPIs & measurable outcomes")
    for kpi in kpis:
        _bullet(doc, kpi)
    _h3(doc, "Typical outcomes (what clients can expect)")
    for outcome in typical_outcomes:
        _bullet(doc, outcome)
    _h3(doc, "Example in practice (simplified)")
    for step in example_steps:
        _num(doc, step)


def _add_value_summary_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Outcome"
    hdr[1].text = "What clients get"
    hdr[2].text = "Who benefits"
    hdr[3].text = "Why it matters"

    rows = [
        (
            "Stock accuracy",
            "A ledger-based inventory engine with controlled execution",
            "Warehouse, operations, finance",
            "Reduces stock discrepancies and write-offs",
        ),
        (
            "Governance",
            "Role-based access and optional approvals before execution",
            "Management, compliance",
            "Prevents unauthorized or accidental stock changes",
        ),
        (
            "Speed at the floor",
            "Barcode/QR scan & lookup for fast identification",
            "Warehouse teams",
            "Faster receiving/picking and fewer manual entry errors",
        ),
        (
            "Financial control",
            "FIFO/Average costing and margin visibility (where enabled)",
            "Finance, leadership",
            "Improves decision-making with consistent inventory value",
        ),
        (
            "Planning",
            "Smart reorder suggestions and stockout prediction",
            "Operations, procurement",
            "Supports service levels and business continuity",
        ),
        (
            "Transparency",
            "Audit logging and traceability",
            "Auditors, management",
            "Creates accountability and supports compliance",
        ),
    ]

    for outcome, what, who, why in rows:
        row = table.add_row().cells
        row[0].text = outcome
        row[1].text = what
        row[2].text = who
        row[3].text = why


def build_doc() -> Document:
    doc = Document()
    _set_normal_style(doc)
    _add_header_footer(doc)
    _cover_page(doc)
    _toc(doc)

    # =========================================================
    # SECTION 1 — EXECUTIVE SUMMARY
    # =========================================================
    _h1(doc, "SECTION 1 — Executive Summary")
    _p(
        doc,
        "The Inventory Management System is an enterprise-ready platform for controlling inventory operations across multiple warehouses with strong governance, auditability, and performance. "
        "It is designed to reduce operational risk, improve stock accuracy, and enable predictable supply planning—without adding complexity for frontline users.",
    )
    _h2(doc, "What the system is")
    _bullet(doc, "A modern web-based inventory and stock control platform for multi-site operations.")
    _bullet(doc, "A single source of truth for on-hand stock, movements, approvals, and reporting.")
    _bullet(doc, "A scalable foundation for integrations and future automation.")

    _h2(doc, "Who it is designed for")
    _bullet(doc, "Warehousing and inventory operations teams")
    _bullet(doc, "Procurement and supply planning")
    _bullet(doc, "Sales operations needing reliable availability")
    _bullet(doc, "Finance and leadership requiring consistent costing visibility")
    _bullet(doc, "Compliance and audit stakeholders")

    _h2(doc, "Key value proposition")
    _bullet(doc, "Higher stock accuracy with controlled execution and full traceability.")
    _bullet(doc, "Fewer preventable losses through approval safeguards and user accountability.")
    _bullet(doc, "Faster daily operations through barcode-enabled workflows and efficient lookup.")
    _bullet(doc, "Smarter replenishment decisions through forecasting and reorder intelligence.")

    _h2(doc, "Enterprise readiness")
    _bullet(doc, "Permission-based access control and customizable roles.")
    _bullet(doc, "Optional approvals for high-impact actions and sensitive workflows.")
    _bullet(doc, "Audit logs designed for accountability and compliance reporting.")
    _bullet(doc, "Performance-aware design suitable for growing catalogs and multi-warehouse environments.")

    _h2(doc, "Competitive positioning")
    _bullet(doc, "Balances ease of use for operators with governance controls demanded by enterprises.")
    _bullet(doc, "Built around accuracy, auditability, and predictable planning—not spreadsheets and guesswork.")
    _bullet(doc, "Designed to scale from a few warehouses to large multi-site operations.")

    _h2(doc, "What success looks like (first 90 days)")
    _bullet(doc, "Users execute standardized receiving, sales confirmation, transfers, and adjustments with minimal training.")
    _bullet(doc, "Stock accuracy improves and reconciliation effort drops measurably.")
    _bullet(doc, "Approval workflows are in place for sensitive operations and exceptions are clearly visible.")
    _bullet(doc, "Low-stock risk and reorder suggestions support proactive procurement.")
    _bullet(doc, "Leadership gains trusted reporting for operational and financial decision-making.")

    _h2(doc, "At-a-glance outcomes")
    _add_value_summary_table(doc)
    doc.add_page_break()

    # =========================================================
    # SECTION 2 — BUSINESS CHALLENGES WE SOLVE
    # =========================================================
    _h1(doc, "SECTION 2 — Business Challenges We Solve")
    _h2(doc, "Manual stock errors")
    _p(
        doc,
        "Manual updates, disconnected spreadsheets, and inconsistent processes create a predictable pattern: stock counts drift, availability becomes unreliable, and teams spend time reconciling instead of operating.",
    )
    _bullet(doc, "Reduce data entry errors by standardizing workflows and adding scan support.")
    _bullet(doc, "Increase confidence in availability across teams and warehouses.")
    _bullet(doc, "Improve training outcomes with consistent steps and fewer ad-hoc workarounds.")
    _bullet(doc, "Reduce dependence on a few ‘power users’ who hold process knowledge.")

    _h2(doc, "Lack of approval control")
    _p(
        doc,
        "Enterprises need controls that prevent accidental or unauthorized stock changes—especially for high-value, regulated, or high-volume items. "
        "Approvals create a clear governance layer without slowing everyday work.",
    )
    _bullet(doc, "Add decision checkpoints for sensitive actions without blocking normal flow.")
    _bullet(doc, "Ensure actions are reviewed and executed consistently.")
    _bullet(doc, "Support segregation of duties (request vs approve) where required.")
    _bullet(doc, "Create a predictable, auditable record of review and outcome.")

    _h2(doc, "Inaccurate costing")
    _p(
        doc,
        "When costing is inconsistent or opaque, organizations lose margin visibility and cannot trust financial reports. "
        "The system supports structured costing approaches that align inventory value with operational reality.",
    )
    _bullet(doc, "Improve confidence in inventory value and cost visibility (where enabled).")
    _bullet(doc, "Reduce manual spreadsheet reconciliations and month-end surprises.")

    _h2(doc, "No forecasting")
    _p(
        doc,
        "Stockouts and overstock are two sides of the same problem: lack of planning intelligence. "
        "Forecasting and reorder suggestions help organizations maintain service levels while controlling working capital.",
    )
    _bullet(doc, "Support proactive replenishment instead of reactive expediting.")
    _bullet(doc, "Reduce overstock by aligning reorder decisions with usage patterns.")

    _h2(doc, "Poor audit visibility")
    _p(
        doc,
        "Without traceability, it becomes hard to understand what changed, who changed it, and why. "
        "Audit visibility is a core requirement for many enterprise clients and regulated environments.",
    )
    _bullet(doc, "Strengthen accountability by linking actions to users and context.")
    _bullet(doc, "Accelerate investigations and exception management.")

    _h2(doc, "Multi-warehouse complexity")
    _p(
        doc,
        "As organizations grow, inventory stops being a single-location problem. Transfers, warehouse-level availability, and controlled movement become essential for accurate operations and customer commitments.",
    )
    _bullet(doc, "Improve warehouse-to-warehouse coordination and visibility.")
    _bullet(doc, "Reduce mis-shipments and misallocation across sites.")

    _h2(doc, "Lack of barcode efficiency")
    _p(
        doc,
        "Barcode and QR workflows reduce picking/receiving time, improve accuracy, and make training easier. "
        "The system supports quick scan/lookup patterns for warehouse execution.",
    )
    _bullet(doc, "Increase throughput while reducing errors in receiving and picking.")
    _bullet(doc, "Make item identification faster across products, batches, and serial-tracked items.")
    doc.add_page_break()

    # =========================================================
    # SECTION 3 — SYSTEM OVERVIEW
    # =========================================================
    _h1(doc, "SECTION 3 — System Overview")
    _p(
        doc,
        "This section provides a high-level overview of the system’s capabilities. It focuses on business outcomes and operational value, rather than internal implementation details.",
    )
    _h2(doc, "Where this system fits best")
    _bullet(doc, "Distribution and logistics operations that need reliable multi-warehouse execution.")
    _bullet(doc, "Retail and wholesale teams with growing catalogs and replenishment needs.")
    _bullet(doc, "Manufacturing support inventory where traceability and controlled adjustments matter.")
    _bullet(doc, "Regulated or high-value inventory requiring batch/serial tracking and audit trails.")
    _bullet(doc, "Organizations transitioning away from spreadsheet-driven inventory control.")

    _h2(doc, "Inventory engine")
    _bullet(doc, "Designed for accuracy: inventory changes follow controlled workflows.")
    _bullet(doc, "Built for traceability: every meaningful action can be audited.")
    _bullet(doc, "Supports high throughput operations without sacrificing governance.")

    _h2(doc, "Multi-warehouse support")
    _bullet(doc, "Warehouse-level availability with clear movement history.")
    _bullet(doc, "Transfers that preserve accountability across sites.")
    _bullet(doc, "Consistent reporting across the network.")

    _h2(doc, "Real-time stock tracking")
    _bullet(doc, "Up-to-date availability after receiving, sales confirmation, adjustments, and transfers.")
    _bullet(doc, "Operational confidence for teams relying on accurate stock positions.")

    _h2(doc, "Approval workflows")
    _bullet(doc, "Optional governance for high-impact actions.")
    _bullet(doc, "Reviewer model: request → review → approve/reject → execution.")
    _bullet(doc, "Idempotent execution: prevents duplicate execution from repeated actions.")

    _h2(doc, "Valuation & costing")
    _bullet(doc, "Supports standard costing approaches used by enterprise clients.")
    _bullet(doc, "Improves margin visibility and financial controls (where enabled).")

    _h2(doc, "Smart reorder system")
    _bullet(doc, "Suggested reorder quantities based on usage and policies.")
    _bullet(doc, "Stockout prediction to protect service levels.")
    _bullet(doc, "Operational continuity: reduce urgent buying and exceptions.")

    _h2(doc, "Barcode & scanning support")
    _bullet(doc, "Scan/lookup to quickly identify products, batches, and serial items.")
    _bullet(doc, "Supports both handheld/camera scanning and USB scanners.")
    doc.add_page_break()

    # =========================================================
    # SECTION 4 — CORE FEATURES
    # =========================================================
    _h1(doc, "SECTION 4 — Core Features")
    _p(
        doc,
        "Core features are organized for enterprise buyers: each capability includes business benefit, operational benefit, and risk mitigation benefit.",
    )
    _p(
        doc,
        "For each feature below, we describe the practical capability, the value it delivers, and the controls that make it suitable for enterprise operations. "
        "This format supports executive stakeholders as well as operational leaders evaluating fit for rollout.",
    )

    _feature_section(
        doc,
        title="1) Inventory Management",
        overview=(
            "Maintain a reliable inventory position by standardizing how stock is received, confirmed, corrected, and transferred. "
            "The system is designed so everyday operations stay fast for users while management retains control and visibility."
        ),
        key_capabilities=[
            "Centralized product catalog with consistent attributes and operational settings",
            "Warehouse-level availability visibility for daily execution",
            "Controlled stock-impacting actions aligned to operational reality",
            "Clear references and context for stock-impacting events (e.g., receiving and confirmation)",
            "Role-based visibility so stakeholders see what they need without broad access",
        ],
        business_benefit="Improve stock accuracy and reduce reconciliation time, enabling better service levels and customer trust.",
        operational_benefit="Streamline day-to-day receiving, confirming sales, and visibility of availability by location.",
        risk_benefit="Reduce losses from uncontrolled changes and improve traceability for investigations and audits.",
        operational_notes=[
            "Designed for frontline speed with standardized steps and clear fields",
            "Balances visibility needs across operations, procurement, sales, and management",
            "Supports governance requirements without forcing unnecessary steps for low-risk actions",
            "Provides consistent definitions for stock changes (receiving, sale confirmation, adjustment, transfer)",
            "Enables clean handoffs between teams through shared references and history",
        ],
        kpis=[
            "Reduction in inventory discrepancy rate (cycle count variance)",
            "Reduction in reconciliation effort (hours per week/month)",
            "Improved fulfillment accuracy (mis-shipments tied to inventory errors)",
            "Improved availability confidence (fewer stock-related backorders)",
            "Faster onboarding time for new warehouse users",
        ],
        typical_outcomes=[
            "Fewer inventory discrepancies and faster issue resolution",
            "Less time spent reconciling spreadsheets across teams",
            "More reliable availability information for operations and sales",
        ],
        example_steps=[
            "Warehouse staff receives goods and records quantities against a warehouse location.",
            "If batch/serial applies, details are captured during the receiving process.",
            "Inventory becomes available for downstream processes and reporting.",
        ],
    )

    _feature_section(
        doc,
        title="2) Multi-Warehouse Operations",
        overview=(
            "Operate multiple sites with clear warehouse-level accountability. The system supports visibility by site, structured transfers, "
            "and reporting that helps leadership understand inventory distribution and constraints."
        ),
        key_capabilities=[
            "Warehouse master data and consistent operational definitions",
            "Availability and movement visibility by warehouse",
            "Controlled transfers that preserve accountability",
            "Site-to-site traceability for investigations and performance review",
            "Support for growing the warehouse network without process breakdown",
        ],
        business_benefit="Support growth across sites without losing control of inventory and fulfillment performance.",
        operational_benefit="Track availability and movements by warehouse; execute controlled transfers.",
        risk_benefit="Reduce mis-shipments and location confusion; maintain clear movement history.",
        operational_notes=[
            "Warehouse-to-warehouse visibility reduces dependency on informal communication",
            "Transfers create accountability for both the sending and receiving locations",
            "Supports standardized naming and reporting across new sites",
            "Enables operational planning by understanding where inventory actually sits",
            "Designed to handle growth in transaction volume without sacrificing traceability",
        ],
        kpis=[
            "Reduction in transfer-related discrepancies",
            "Improved warehouse-level stock accuracy",
            "Reduction in emergency transfers and last-minute reallocations",
            "Improved service level by correct inventory positioning",
            "Faster resolution of cross-warehouse issues",
        ],
        typical_outcomes=[
            "Fewer site-level surprises and more predictable fulfillment performance",
            "Improved decision-making for where to position stock",
            "Clear accountability for transfers and warehouse-level inventory health",
        ],
        example_steps=[
            "A manager identifies imbalance: Warehouse A is overstocked; Warehouse B is at risk of stockout.",
            "A transfer is created to move the required quantity from A to B.",
            "Both warehouses show the movement history and updated availability.",
        ],
    )

    _feature_section(
        doc,
        title="3) Purchase & Sales Management",
        overview=(
            "Purchases and sales are aligned with how enterprises operate: inventory impact happens at the correct operational moment. "
            "This reduces confusion and ensures reporting reflects real execution."
        ),
        key_capabilities=[
            "Receiving increases stock when goods are confirmed as received",
            "Sales reduce stock when a sale is confirmed/executed (or after approval when required)",
            "Reference numbers and context for operational traceability",
            "Warehouse-level execution aligned to fulfillment reality",
            "Improved reporting clarity across purchase and sales activity",
        ],
        business_benefit="Better order fulfillment reliability with dependable availability; fewer surprises for sales operations.",
        operational_benefit="Receiving increases stock; confirming sales decreases stock with clear reference tracking.",
        risk_benefit="Avoid premature stock impacts; reduce errors caused by partial or unverified events.",
        operational_notes=[
            "Separates ‘planned’ events from ‘executed’ events to avoid confusion in reporting",
            "Supports clean handoffs between receiving teams and downstream fulfillment",
            "Improves alignment between sales commitments and actual availability",
            "Provides consistent references for dispute resolution and exception analysis",
            "Supports optional approvals where enterprise governance requires review",
        ],
        kpis=[
            "Reduction in stock-related order exceptions",
            "Improved on-time fulfillment due to accurate availability signals",
            "Reduction in manual corrections caused by early/incorrect stock impacts",
            "Improved accuracy of sales and receiving reporting timelines",
            "Reduction in customer escalations tied to inventory mismatch",
        ],
        typical_outcomes=[
            "Reduced disputes between operations and sales on availability",
            "Clearer operational reporting and better customer commitments",
            "Lower exception rate caused by early or incorrect stock updates",
        ],
        example_steps=[
            "Goods arrive and are received into a warehouse with validated quantities.",
            "A sale is confirmed when it is ready to be executed and shipped/fulfilled.",
            "Reports reflect true receiving and true sales execution timing.",
        ],
    )

    _feature_section(
        doc,
        title="4) Stock Adjustments & Transfers",
        overview=(
            "Adjustments and transfers handle real-world exceptions: damage, cycle count corrections, and redistribution of stock. "
            "The system makes these events visible and auditable, protecting enterprises from silent drift."
        ),
        key_capabilities=[
            "Controlled stock adjustments with reasons and accountability",
            "Transfers between warehouses with end-to-end traceability",
            "Visibility into exception patterns for continuous improvement",
            "Policy controls to prevent high-risk actions without review",
            "Consistent reporting of corrections and movements",
        ],
        business_benefit="Reduce shrinkage impact by quickly correcting inventory and keeping operations stable.",
        operational_benefit="Standardized adjustment and transfer workflows with consistent reporting.",
        risk_benefit="Minimize unauthorized edits; ensure transfers are fully accounted for on both sides.",
        operational_notes=[
            "Adjustments are treated as controlled exceptions, not silent edits",
            "Transfers maintain end-to-end accountability and visible history",
            "Supports process improvement by surfacing exception patterns over time",
            "Helps leadership separate operational variance from governance issues",
            "Designed to support high-volume operations with clear oversight",
        ],
        kpis=[
            "Reduction in unexplained adjustment volume (trend over time)",
            "Time-to-resolution for stock discrepancies",
            "Reduction in negative events (e.g., missing stock during fulfillment)",
            "Improved audit readiness for correction activity",
            "Improved accuracy of transfer execution between warehouses",
        ],
        typical_outcomes=[
            "Improved cycle count accuracy and faster correction handling",
            "Fewer unexplained variances at month-end",
            "Clear visibility into shrinkage and operational exceptions",
        ],
        example_steps=[
            "A cycle count reveals an overage/shortage for an item in a warehouse location.",
            "A controlled adjustment is recorded with a reason (e.g., count correction, damage).",
            "Managers can review patterns and take corrective process actions.",
        ],
    )

    _feature_section(
        doc,
        title="5) Batch & Serial Tracking",
        overview=(
            "For regulated, high-value, or expiry-sensitive inventory, batch and serial tracking enable deep traceability. "
            "This supports compliance, recall readiness, and stronger customer confidence."
        ),
        key_capabilities=[
            "Batch visibility for expiry or lot-controlled items",
            "Serial tracking for unique items requiring individual traceability",
            "Warehouse-level traceability for movement and status changes",
            "Faster investigations during issues or customer inquiries",
            "Reduced risk of shipping incorrect or non-compliant stock",
        ],
        business_benefit="Enable traceability that supports compliance, recall readiness, and customer requirements.",
        operational_benefit="Track inventory at batch and serial level where needed; maintain visibility across warehouses.",
        risk_benefit="Reduce the risk of shipping wrong items; support investigations with precise traceability.",
        operational_notes=[
            "Supports organizations with traceability obligations (expiry, lot control, regulated goods)",
            "Improves customer confidence through more precise inventory answers",
            "Enables faster issue handling when exceptions arise (returns, damage, recalls)",
            "Reduces training complexity by embedding traceability into standard workflows",
            "Maintains warehouse-level visibility so operations stay predictable",
        ],
        kpis=[
            "Reduction in traceability-related fulfillment errors",
            "Time-to-answer for customer traceability inquiries",
            "Reduction in expiry-related losses (where applicable)",
            "Improved compliance readiness (audit response time)",
            "Reduction in manual traceability spreadsheets or side systems",
        ],
        typical_outcomes=[
            "Stronger compliance posture for batch/serial-sensitive products",
            "Reduced customer escalations through faster traceability answers",
            "More controlled handling of expiry and regulated inventory",
        ],
        example_steps=[
            "Receiving captures batch information (and serials where required).",
            "Operations can locate the correct batch/serial stock for fulfillment.",
            "Audit trails support recall readiness and investigations if needed.",
        ],
    )

    _feature_section(
        doc,
        title="6) Financial Valuation (FIFO/Average)",
        overview=(
            "Enterprises need inventory costing approaches that align with governance and reporting. "
            "The system supports structured valuation methods and improves financial visibility where enabled."
        ),
        key_capabilities=[
            "Support for FIFO and average costing approaches",
            "Consistent cost handling to reduce reporting friction",
            "Improved visibility into cost and margin signals (where enabled)",
            "Clear auditability of valuation-related outcomes",
            "Controls to restrict financial visibility based on role/permission",
        ],
        business_benefit="Improve margin visibility and financial planning with standardized costing.",
        operational_benefit="Consistent cost handling reduces downstream reporting friction and rework.",
        risk_benefit="Reduce disputes and audit risk by applying consistent costing logic.",
        operational_notes=[
            "Supports finance stakeholders with consistent valuation visibility (where enabled)",
            "Reduces reliance on manual costing spreadsheets and ad-hoc adjustments",
            "Supports governance by restricting financial visibility when required",
            "Aligns operational execution with financial outcomes and accountability",
            "Improves consistency across warehouses and product categories",
        ],
        kpis=[
            "Reduction in manual financial reconciliation effort",
            "Improved margin reporting consistency over time",
            "Reduction in valuation disputes or corrections",
            "Improved timeliness of financial reporting related to inventory",
            "Reduction in cost anomalies flagged by leadership",
        ],
        typical_outcomes=[
            "More reliable inventory valuation visibility for finance stakeholders",
            "Reduced manual adjustments and reconciliation effort",
            "Improved decision-making for pricing and procurement",
        ],
        example_steps=[
            "A client selects FIFO or average costing based on their reporting requirements.",
            "As sales are executed, cost-of-goods visibility aligns with the configured method.",
            "Leadership uses consistent signals for margin and financial performance review.",
        ],
    )

    _feature_section(
        doc,
        title="7) Smart Reorder & Forecasting",
        overview=(
            "The system provides practical replenishment guidance: low stock signals, suggested reorder quantities, and stockout prediction. "
            "This improves service levels and reduces expensive reactive procurement."
        ),
        key_capabilities=[
            "Low stock visibility and prioritized replenishment signals",
            "Suggested reorder quantities aligned to usage and policy",
            "Stockout risk prediction to protect service levels",
            "Warehouse-level planning visibility",
            "Support for ongoing tuning as the business learns and scales",
        ],
        business_benefit="Support continuity and service levels; reduce emergency buying and expedite costs.",
        operational_benefit="Surface low stock, suggested reorder quantities, and predicted stockout signals.",
        risk_benefit="Reduce operational disruption from stockouts and reduce cash tied in excess inventory.",
        operational_notes=[
            "Designed for practical planning: alerts and suggestions are action-oriented",
            "Supports warehouse-level planning rather than one-size-fits-all replenishment",
            "Encourages policy-driven replenishment (lead time, safety stock, min/max levels)",
            "Helps leadership balance working capital with service levels",
            "Enables ongoing tuning as demand patterns evolve",
        ],
        kpis=[
            "Reduction in stockout incidents on priority items",
            "Improved fill rate / service level",
            "Reduction in expediting and urgent procurement costs",
            "Reduction in excess inventory on slow-moving items",
            "Improved planning cycle time (time spent deciding what to reorder)",
        ],
        typical_outcomes=[
            "Higher fill rates and fewer urgent procurement events",
            "Lower overstock on slow-moving items and improved working capital",
            "More predictable operations across warehouses",
        ],
        example_steps=[
            "A planner reviews low stock and predicted stockout risk by warehouse.",
            "The system provides a suggested reorder quantity aligned to policy.",
            "Procurement prioritizes actions based on service level risk and business importance.",
        ],
    )

    _feature_section(
        doc,
        title="8) Approval Workflow Engine",
        overview=(
            "Approvals add governance for enterprises that require control over sensitive workflows. "
            "Rather than blocking operations, the system routes actions to review when needed and executes them once approved."
        ),
        key_capabilities=[
            "Configurable approval policies for key workflows",
            "Clear request lifecycle: pending → approved/rejected",
            "Reviewer accountability and decision tracking",
            "Execution safeguards to prevent double execution",
            "Visibility into pending work to prevent operational bottlenecks",
        ],
        business_benefit="Strengthen controls and reduce costly mistakes in high-volume operations.",
        operational_benefit="Request/review/approve execution model keeps teams aligned and accountable.",
        risk_benefit="Prevents unauthorized actions; ensures sensitive workflows are reviewed and traceable.",
        operational_notes=[
            "Supports segregation of duties and enterprise governance requirements",
            "Prevents sensitive actions from being executed without review",
            "Creates a clear queue of pending work to reduce bottlenecks",
            "Ensures decisions are traceable, including reviewer identity and timing",
            "Designed to prevent duplicate execution in high-volume environments",
        ],
        kpis=[
            "Reduction in high-impact errors (large adjustments, sensitive transfers)",
            "Approval cycle time (request → decision)",
            "Reduction in policy violations or unauthorized changes",
            "Improved audit readiness for approval-required workflows",
            "Reduction in rework due to incorrect or unreviewed actions",
        ],
        typical_outcomes=[
            "Lower frequency of high-impact errors (e.g., large adjustments, sensitive transfers)",
            "Improved governance confidence for auditors and leadership",
            "Clear review queues that balance control and operational speed",
        ],
        example_steps=[
            "A staff member requests a sensitive adjustment (or sale confirmation) that requires review.",
            "A manager reviews and approves/rejects with optional notes.",
            "If approved, the system executes the action once and records the full audit trail.",
        ],
    )

    _feature_section(
        doc,
        title="9) Barcode & QR Scanning",
        overview=(
            "Barcode/QR scan-first workflows reduce typing, speed up receiving and picking, and improve accuracy. "
            "The system supports quick scan/lookup to find products, batches, or serial items."
        ),
        key_capabilities=[
            "Fast scan/lookup to identify products and stock context",
            "Support for camera-based scanning and USB scanners",
            "Improved usability for frontline teams",
            "Reduced manual entry and training time",
            "Better accuracy for high-throughput workflows",
        ],
        business_benefit="Faster receiving, picking, and verification improves productivity and customer outcomes.",
        operational_benefit="Scan/lookup reduces manual typing and training time; supports varied devices.",
        risk_benefit="Reduces incorrect item selection and data entry mistakes.",
        operational_notes=[
            "Supports high-throughput warehouse workflows where typing is a bottleneck",
            "Reduces training time by simplifying identification and lookup",
            "Improves accuracy in environments with look-alike products or complex SKUs",
            "Supports multiple scanning approaches (camera or USB scanner) to match client environments",
            "Improves operational confidence by returning consistent item identification results",
        ],
        kpis=[
            "Reduction in picking/receiving cycle time",
            "Reduction in item identification errors",
            "Reduction in manual typing-related data entry errors",
            "Improved throughput per user/hour in high-volume workflows",
            "Reduction in training time to reach target productivity",
        ],
        typical_outcomes=[
            "Reduced receiving/picking cycle time",
            "Lower error rate from manual entry and incorrect item selection",
            "Faster onboarding for new warehouse staff",
        ],
        example_steps=[
            "A user scans an item code using a USB scanner or camera.",
            "The system returns the matching product/batch/serial result with key details.",
            "The user proceeds with the next operation with fewer mistakes and less rework.",
        ],
    )

    _feature_section(
        doc,
        title="10) Audit & Compliance Logging",
        overview=(
            "Audit visibility supports enterprise governance: who did what, when, and with what outcome. "
            "The system is designed to support investigations, compliance needs, and operational accountability."
        ),
        key_capabilities=[
            "Audit trails for high-impact operational actions and governance decisions",
            "Visibility for compliance and management stakeholders",
            "Support for investigations and exception resolution",
            "Approval decision traceability and reviewer accountability",
            "Operational transparency without slowing daily work",
        ],
        business_benefit="Improves governance confidence and supports compliance requirements.",
        operational_benefit="Simplifies investigations by linking actions to users, time, and context.",
        risk_benefit="Reduces fraud risk and supports audits with reliable evidence trails.",
        operational_notes=[
            "Supports internal controls by making key actions visible and reviewable",
            "Helps reduce reliance on informal knowledge and ad-hoc investigations",
            "Improves accountability across teams and warehouses",
            "Enables governance review routines (exceptions, adjustments, approvals)",
            "Supports audit stakeholders with consistent, predictable evidence trails",
        ],
        kpis=[
            "Time-to-resolution for discrepancies and exceptions",
            "Reduction in repeat exceptions (trend over time)",
            "Audit response time improvements (evidence retrieval speed)",
            "Reduction in unauthorized or unexplained activity",
            "Improved management review cadence and coverage",
        ],
        typical_outcomes=[
            "Faster resolution of exceptions and discrepancies",
            "Stronger compliance posture for regulated environments",
            "Higher management confidence in operational governance",
        ],
        example_steps=[
            "A discrepancy is detected in inventory reporting.",
            "Teams review the audit trail to identify actions and responsible stakeholders.",
            "Corrective action is applied and documented with traceability intact.",
        ],
    )

    _feature_section(
        doc,
        title="11) Reporting & Analytics",
        overview=(
            "Enterprise teams need reporting that is actionable, trusted, and role-appropriate. "
            "The system provides dashboards and reports for operational execution, management review, and planning."
        ),
        key_capabilities=[
            "Role-scoped dashboards for daily execution and leadership visibility",
            "Operational reports across inventory, movements, purchases, sales, and audit",
            "Filters for warehouse, category, and time range",
            "Export-ready outputs for downstream workflows where permitted",
            "Signals for low stock and planning priorities",
        ],
        business_benefit="Create shared visibility across leadership, operations, and finance.",
        operational_benefit="Role-scoped dashboards and reports for daily execution and planning.",
        risk_benefit="Detect anomalies earlier; reduce surprises in operations and financial outcomes.",
        operational_notes=[
            "Designed to be actionable: reports support decisions, not just record-keeping",
            "Role-based visibility prevents overexposure of sensitive information",
            "Supports operational management routines (daily/weekly reviews)",
            "Enables exception-focused management rather than reactive firefighting",
            "Improves alignment across teams with shared, trusted data",
        ],
        kpis=[
            "Improved decision cycle time (how fast teams act on insights)",
            "Reduction in recurring exceptions identified via reporting",
            "Improved service levels through early risk visibility",
            "Reduction in time spent creating manual reports",
            "Improved management confidence in operational data quality",
        ],
        typical_outcomes=[
            "More predictable operations through shared visibility and aligned priorities",
            "Earlier detection of anomalies and exception trends",
            "Improved decision-making through trusted reporting",
        ],
        example_steps=[
            "A manager reviews dashboards for low stock risk, movement trends, and exceptions.",
            "Reports are filtered by warehouse and timeframe to support action planning.",
            "Teams export or share key summaries for downstream review where required.",
        ],
    )
    doc.add_page_break()

    # =========================================================
    # SECTION 5 — ROLE-BASED ACCESS & GOVERNANCE
    # =========================================================
    _h1(doc, "SECTION 5 — Role-Based Access & Governance")
    _p(
        doc,
        "Enterprise clients need to ensure the right people can do the right actions—and only those actions. The system uses permission-based access to support both standard and custom roles.",
    )

    _h2(doc, "Standard roles")
    _h3(doc, "Admin")
    _bullet(doc, "Full access across all modules and configuration controls.")
    _bullet(doc, "Responsible for onboarding, governance, and system configuration.")
    _h3(doc, "Manager")
    _bullet(doc, "Operational leadership access for inventory execution and reporting.")
    _bullet(doc, "May act as reviewer for approval workflows.")
    _h3(doc, "Staff")
    _bullet(doc, "Frontline operational execution (as enabled by policy and permissions).")
    _bullet(doc, "Designed for warehouse users with clear, safe workflows.")
    _h3(doc, "Viewer")
    _bullet(doc, "Read-only visibility for stakeholders who need insight without write capability.")

    _h2(doc, "Custom roles")
    _bullet(doc, "Create roles aligned to job function (e.g., procurement, warehouse lead, auditor).")
    _bullet(doc, "Assign specific permissions to ensure least-privilege access.")

    _h2(doc, "Permission-based architecture")
    _bullet(doc, "Permissions enable precise governance beyond role names.")
    _bullet(doc, "Supports segregation of duties (e.g., request vs approve).")

    _h2(doc, "Approval safeguards")
    _bullet(doc, "Sensitive actions can be configured to require review before execution.")
    _bullet(doc, "Review actions are auditable, including outcome and reviewer identity.")

    _h2(doc, "Financial visibility controls")
    _bullet(doc, "Financial metrics and valuation visibility can be restricted by permission.")
    _bullet(doc, "Supports clients who separate operational access from financial access.")
    doc.add_page_break()

    # =========================================================
    # SECTION 6 — OPERATIONAL WORKFLOWS (CLIENT-FRIENDLY)
    # =========================================================
    _h1(doc, "SECTION 6 — Operational Workflows (Client-Friendly)")
    _p(
        doc,
        "Below are the primary workflows described in business terms. The system is designed to be predictable for frontline teams and controllable for enterprise governance.",
    )

    _h2(doc, "Purchase to Stock flow")
    _num(doc, "Receive stock for a warehouse against a supplier shipment or reference.")
    _num(doc, "Validate quantities, batches/serials (if applicable), and confirm receipt.")
    _num(doc, "Stock becomes available immediately for downstream operations and reporting.")

    _h2(doc, "Sale to COGS flow")
    _num(doc, "Create or record a sale order for a warehouse.")
    _num(doc, "Confirm the sale when it is ready to be executed (or route to approval if policy requires).")
    _num(doc, "Stock is reduced and the transaction is traceable for reporting and costing visibility.")

    _h2(doc, "Transfer between warehouses")
    _num(doc, "Select source warehouse, destination warehouse, and quantity.")
    _num(doc, "Confirm transfer; inventory moves out of the source and into the destination with a single traceable transfer record.")
    _num(doc, "Both warehouses reflect the movement and history for accountability.")

    _h2(doc, "Adjustment & correction flow")
    _num(doc, "When discrepancies occur, apply an adjustment with a reason (e.g., damage, count correction).")
    _num(doc, "The system records what changed and who performed the correction.")
    _num(doc, "Reporting retains visibility of adjustments for audits and improvement actions.")

    _h2(doc, "Approval lifecycle")
    _num(doc, "A user requests an action that is configured as approval-required.")
    _num(doc, "A reviewer approves or rejects with optional notes.")
    _num(doc, "On approval, the action executes once and becomes part of the traceable operational history.")

    _h2(doc, "Reorder alert handling")
    _num(doc, "The system monitors stock levels and identifies low-stock or stockout risk.")
    _num(doc, "Suggested reorder quantity is provided based on policy and usage trends.")
    _num(doc, "Teams can action reorder decisions with clear reasoning and prioritization.")

    _h2(doc, "Scan & lookup process")
    _num(doc, "A user scans a barcode/QR code or enters a code manually.")
    _num(doc, "The system returns the most relevant match (product, batch, or serial) and key availability details.")
    _num(doc, "Users proceed with receiving, picking, or verification with fewer errors and faster execution.")
    doc.add_page_break()

    # =========================================================
    # SECTION 7 — FINANCIAL CONTROL & COSTING
    # =========================================================
    _h1(doc, "SECTION 7 — Financial Control & Costing")
    _p(
        doc,
        "Inventory value and cost visibility are essential for enterprise management. The system supports common valuation approaches used in practice and provides consistent, auditable outcomes.",
    )
    _h2(doc, "FIFO valuation")
    _bullet(doc, "Supports cost layering and consistent valuation outcomes.")
    _bullet(doc, "Useful for organizations that need disciplined cost tracking over time.")
    _h2(doc, "Average costing")
    _bullet(doc, "Provides stable cost behavior and simplified valuation for high-volume items.")
    _bullet(doc, "Suitable for businesses with frequent replenishment and consistent purchasing patterns.")
    _h2(doc, "Margin visibility")
    _bullet(doc, "Supports informed decisions on pricing, procurement, and operational trade-offs.")
    _bullet(doc, "Helps identify margin leakage and cost anomalies earlier.")
    _h2(doc, "Financial reporting benefits")
    _bullet(doc, "Consistent cost handling reduces reconciliation overhead.")
    _bullet(doc, "Improves confidence for leadership and audit stakeholders.")
    doc.add_page_break()

    # =========================================================
    # SECTION 8 — FORECASTING & INTELLIGENCE
    # =========================================================
    _h1(doc, "SECTION 8 — Forecasting & Intelligence")
    _p(
        doc,
        "Forecasting capabilities reduce operational disruption and support predictable service levels. The system provides practical signals that teams can act on—without needing data science resources.",
    )
    _h2(doc, "Low stock alerts")
    _bullet(doc, "Identify low stock before it becomes a service issue.")
    _bullet(doc, "Focus on what matters: high-impact products and key locations.")
    _h2(doc, "Suggested reorder quantity")
    _bullet(doc, "Recommendations align with usage trends and configured policies.")
    _bullet(doc, "Helps maintain service levels while controlling inventory investment.")
    _h2(doc, "Stockout prediction")
    _bullet(doc, "Early warning signals protect continuity and customer commitments.")
    _bullet(doc, "Supports proactive procurement and replenishment scheduling.")
    _h2(doc, "Business continuity advantage")
    _bullet(doc, "Reduce urgent purchasing and expediting costs.")
    _bullet(doc, "Maintain stability across warehouses and distribution points.")
    doc.add_page_break()

    # =========================================================
    # SECTION 9 — COMPLIANCE & AUDIT
    # =========================================================
    _h1(doc, "SECTION 9 — Compliance & Audit")
    _p(
        doc,
        "Traceability and accountability are built into the operational design. The system provides a dependable evidence trail for review and audits—while keeping everyday workflows efficient.",
    )
    _h2(doc, "Full traceability")
    _bullet(doc, "Clear history of stock changes by product and warehouse.")
    _bullet(doc, "Context for why changes happened (references, reasons, approvals).")
    _h2(doc, "Immutable ledger mindset")
    _bullet(doc, "Operational history is preserved so investigations are possible and consistent.")
    _h2(doc, "Approval audit trails")
    _bullet(doc, "Request, reviewer, decision, and execution are traceable.")
    _bullet(doc, "Supports segregation of duties and governance evidence.")
    _h2(doc, "User accountability")
    _bullet(doc, "Actions are tied to user identity, role permissions, and timestamps.")
    doc.add_page_break()

    # =========================================================
    # SECTION 10 — SCALABILITY & ARCHITECTURE (HIGH-LEVEL)
    # =========================================================
    _h1(doc, "SECTION 10 — Scalability & Architecture (High-Level)")
    _p(
        doc,
        "The system is built on a modern, scalable web architecture suitable for enterprise environments. This section is intentionally high-level and client-friendly.",
    )
    _h2(doc, "Modern web architecture")
    _bullet(doc, "Web-based access for distributed teams and multi-site operations.")
    _bullet(doc, "Designed for reliability and clear operational flows.")
    _h2(doc, "Secure role-based access")
    _bullet(doc, "Permission-based model to support least-privilege access.")
    _bullet(doc, "Supports customization for client-specific governance models.")
    _h2(doc, "Performance optimized")
    _bullet(doc, "Designed for responsive daily operations and reporting workloads.")
    _bullet(doc, "Supports growth in catalog size and warehouse activity.")
    _h2(doc, "Extensible for integrations")
    _bullet(doc, "Designed to be integration-ready for future needs (ERP, accounting, e-commerce, BI).")
    _bullet(doc, "Approach supports staged rollout of integrations to reduce implementation risk.")
    _h2(doc, "API ready (future integration)")
    _bullet(doc, "Supports future integration strategy without redesigning core workflows.")
    doc.add_page_break()

    # =========================================================
    # SECTION 11 — IMPLEMENTATION APPROACH
    # =========================================================
    _h1(doc, "SECTION 11 — Implementation Approach")
    _p(
        doc,
        "A successful rollout balances speed with risk control. The approach below supports enterprise clients through configuration, migration, training, and go-live readiness.",
    )
    _h2(doc, "Suggested implementation phases (example)")
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Phase"
    hdr[1].text = "Focus"
    hdr[2].text = "Typical deliverables"
    phases = [
        ("Phase 1", "Discovery & configuration", "Process mapping, warehouse setup, role/permission model, approval policy design"),
        ("Phase 2", "Data onboarding", "Product catalog import, opening stock validation, pilot warehouse onboarding"),
        ("Phase 3", "Pilot rollout", "Role-based training, controlled go-live for a subset of users/sites"),
        ("Phase 4", "Enterprise rollout", "Full site rollout, governance tuning, reporting baselines and KPIs"),
        ("Phase 5", "Optimization", "Forecast tuning, exception management, continuous improvement cadence"),
    ]
    for phase, focus, deliverables in phases:
        row = table.add_row().cells
        row[0].text = phase
        row[1].text = focus
        row[2].text = deliverables
    _h2(doc, "Setup & configuration")
    _bullet(doc, "Define warehouses, product catalog structure, and operational policies.")
    _bullet(doc, "Configure roles and permissions aligned to organizational governance.")
    _bullet(doc, "Enable approvals for workflows where required.")
    _h2(doc, "Data migration")
    _bullet(doc, "Import products, initial stock, and supporting master data.")
    _bullet(doc, "Validate counts and reconcile variances before go-live.")
    _h2(doc, "User onboarding")
    _bullet(doc, "Create user accounts and assign roles based on job function.")
    _bullet(doc, "Pilot with a small user group to validate workflows.")
    _h2(doc, "Training")
    _bullet(doc, "Role-based training paths: staff, managers, approvers, admins.")
    _bullet(doc, "Operational job aids for receiving, scanning, transfers, and adjustments.")
    _h2(doc, "Go-live checklist")
    _bullet(doc, "Confirm governance settings: approvals, permissions, and financial visibility.")
    _bullet(doc, "Validate stock positions and reporting baselines.")
    _bullet(doc, "Define support process and escalation paths for early go-live weeks.")
    _bullet(doc, "Confirm operational responsibilities (who receives, who approves, who adjusts, who audits).")
    _bullet(doc, "Define a regular cadence for reviewing exceptions, low stock risk, and process improvements.")
    doc.add_page_break()

    # =========================================================
    # SECTION 12 — WHY THIS SYSTEM
    # =========================================================
    _h1(doc, "SECTION 12 — Why This System")
    _bullet(doc, "Enterprise-grade controls without enterprise-level complexity for frontline teams.")
    _bullet(doc, "Scalable multi-warehouse design that grows with operations.")
    _bullet(doc, "Intelligent forecasting and reorder guidance for better continuity.")
    _bullet(doc, "Operational efficiency through scan-enabled workflows and standardized execution.")
    _bullet(doc, "Risk reduction through approvals, traceability, and user accountability.")
    _bullet(doc, "Cost accuracy and margin visibility through structured valuation approaches (where enabled).")
    _bullet(doc, "A future-proof platform designed for integration and extension.")
    doc.add_page_break()

    # =========================================================
    # SECTION 13 — FUTURE ROADMAP (OPTIONAL ADD-ON)
    # =========================================================
    _h1(doc, "SECTION 13 — Future Roadmap (Optional Add-On)")
    _p(
        doc,
        "Roadmap items can be prioritized based on client needs and rollout maturity. The system is designed to support staged expansion without disrupting core operations.",
    )
    _h2(doc, "API integrations")
    _bullet(doc, "ERP/accounting integration for financial automation.")
    _bullet(doc, "E-commerce and order management integration to synchronize demand and fulfillment.")
    _h2(doc, "Advanced analytics")
    _bullet(doc, "Executive dashboards and operational KPIs by site, product category, and time period.")
    _bullet(doc, "Exception analytics for shrinkage, adjustments, and policy violations.")
    _h2(doc, "AI forecasting")
    _bullet(doc, "Enhanced demand forecasting using seasonality, promotions, and lead time variability.")
    _bullet(doc, "Automated recommendations for reorder and inventory positioning across warehouses.")
    _h2(doc, "Mobile-first warehouse tools")
    _bullet(doc, "Optimized mobile workflows for receiving, cycle counts, picking, and verification.")
    _bullet(doc, "Device-native scanning and offline-friendly execution for constrained environments.")
    doc.add_page_break()

    # =========================================================
    # SECTION 14 — CONCLUSION
    # =========================================================
    _h1(doc, "SECTION 14 — Conclusion")
    _p(
        doc,
        "The Inventory Management System provides enterprise clients with the controls and visibility required to run multi-warehouse inventory operations with confidence. "
        "It combines operational speed (scan-enabled workflows), governance (permissions and approvals), financial discipline (valuation support), and intelligence (reorder and stockout prediction). "
        "The result is a scalable, future-proof platform that reduces risk, improves accuracy, and strengthens decision-making.",
    )
    _p(
        doc,
        "We welcome the opportunity to tailor the rollout plan and governance model to your organization’s operational realities and compliance needs.",
    )

    return doc


def _save_doc(doc: Document) -> str:
    for path in (OUTPUT_PATH, FALLBACK_OUTPUT_PATH):
        try:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            doc.save(path)
            return path
        except Exception:
            continue
    raise RuntimeError("Unable to save the document to the primary or fallback output path.")


def main() -> None:
    doc = build_doc()
    saved_to = _save_doc(doc)
    print(f"Saved: {saved_to}")
    if saved_to != OUTPUT_PATH:
        print(f"Primary path not available on this machine: {OUTPUT_PATH}")
        print("If you need the exact /mnt/data path, run this script in an environment where /mnt/data exists.")


if __name__ == "__main__":
    main()

